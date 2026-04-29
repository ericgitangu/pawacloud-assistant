# SPDX-License-Identifier: MIT
"""One-shot fixture generator for document parser tests.

Run once: `python backend/scripts/make-fixtures.py`
Commits the resulting binaries to backend/tests/fixtures/.
"""

from pathlib import Path

from docx import Document


def make_docx(out: Path) -> None:
    doc = Document()
    doc.add_heading("Sample Document", level=1)
    doc.add_paragraph("This is the first paragraph used by upload tests.")
    doc.add_paragraph("Second paragraph with some content for chunking checks.")
    doc.add_heading("Subsection", level=2)
    doc.add_paragraph("Final paragraph.")
    doc.save(out)


def make_pdf(out: Path) -> None:
    """Hand-written single-page PDF with a literal text stream — no extra deps."""
    raw = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R"
        b"/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 76>>stream\n"
        b"BT /F1 12 Tf 72 720 Td (Sample PDF for upload tests.) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n"
        b"0000000009 00000 n \n0000000055 00000 n \n0000000101 00000 n \n"
        b"0000000196 00000 n \n0000000316 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n378\n%%EOF\n"
    )
    out.write_bytes(raw)


if __name__ == "__main__":
    here = Path(__file__).resolve().parent.parent
    fixtures = here / "tests" / "fixtures"
    fixtures.mkdir(parents=True, exist_ok=True)
    make_docx(fixtures / "sample.docx")
    make_pdf(fixtures / "sample.pdf")

    # post-generation parse check — fixtures are no use if they don't load
    Document(fixtures / "sample.docx")
    from pypdf import PdfReader

    reader = PdfReader(str(fixtures / "sample.pdf"))
    extracted = reader.pages[0].extract_text()
    assert "Sample PDF for upload tests" in extracted, "PDF fixture failed parse check"

    print(f"wrote {fixtures}/sample.docx + sample.pdf (verified)")
