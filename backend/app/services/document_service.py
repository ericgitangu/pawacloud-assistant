# SPDX-License-Identifier: MIT
"""Document parsing pipeline — docx + pdf text extraction with normalization."""

import io
import json
import logging
from dataclasses import dataclass
from typing import AsyncGenerator
from uuid import UUID

from docx import Document
from pypdf import PdfReader

from app.core.events import EventCode, emit_event
from app.services.document_processing import (
    chunk_markdown,
    detect_script,
    normalize_document_text,
)
from app.services.text_processing import estimate_tokens, sanitize_input

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"
JPEG_MIME = "image/jpeg"
PNG_MIME = "image/png"
IMAGE_MIMES = {JPEG_MIME, PNG_MIME}


@dataclass
class ParsedDocument:
    text: str
    page_count: int
    char_count: int
    source_lang: str | None
    parse_method: str
    warnings: list[dict]


def parse_document(raw: bytes, filename: str, mime: str) -> ParsedDocument:
    if mime == DOCX_MIME or filename.lower().endswith(".docx"):
        return _parse_docx(raw)
    if mime == PDF_MIME or filename.lower().endswith(".pdf"):
        return _parse_pdf(raw)
    if mime in IMAGE_MIMES or filename.lower().endswith((".jpg", ".jpeg", ".png")):
        return _parse_image(raw, mime)
    raise ValueError(f"unsupported mime: {mime}")


def _parse_docx(raw: bytes) -> ParsedDocument:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    raw_text = "\n\n".join(p for p in parts if p)
    text = normalize_document_text(sanitize_input(raw_text))
    return ParsedDocument(
        text=text,
        page_count=max(1, _approx_page_count(text)),
        char_count=len(text),
        source_lang=detect_script(text, 4096),
        parse_method="docx",
        warnings=[],
    )


def _vision_client():
    """Build a Vision client. Uses GOOGLE_VISION_KEY_PATH when set (local dev);
    falls through to Application Default Credentials on Cloud Run / GCE where
    the runtime service account is the credential source."""
    from app.core.config import get_settings
    from google.cloud import vision

    settings = get_settings()
    if settings.GOOGLE_VISION_KEY_PATH:
        return vision.ImageAnnotatorClient.from_service_account_json(
            settings.GOOGLE_VISION_KEY_PATH
        )
    return vision.ImageAnnotatorClient()


def _try_vision_ocr(raw_pdf: bytes) -> str | None:
    """Send the PDF to Google Cloud Vision for OCR.

    Returns the extracted text on success, None on any failure (missing
    credentials, quota exceeded, network error, malformed PDF). Errors
    are logged but never raised — caller falls back to detect-and-warn.
    """
    try:
        # lazy import — only paid when we actually call this path
        from google.cloud import vision

        client = _vision_client()
        request = vision.AnnotateFileRequest(
            input_config=vision.InputConfig(
                content=raw_pdf,
                mime_type="application/pdf",
            ),
            features=[
                vision.Feature(type_=vision.Feature.Type.DOCUMENT_TEXT_DETECTION)
            ],
        )
        resp = client.batch_annotate_files(requests=[request])
        if not resp.responses:
            return None
        pages_text: list[str] = []
        for file_resp in resp.responses:
            for page_resp in file_resp.responses:
                if page_resp.full_text_annotation:
                    pages_text.append(page_resp.full_text_annotation.text)
        return "\n\n".join(pages_text) if pages_text else None
    except Exception as exc:
        logger.warning("vision ocr failed: %s", exc)
        return None


def _parse_pdf(raw: bytes) -> ParsedDocument:
    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")

    raw_text = "\n\n".join(p for p in parts if p.strip())
    text = normalize_document_text(sanitize_input(raw_text))
    page_count = len(reader.pages)

    parse_method = "pdf-text"
    warnings: list[dict] = []

    # OCR fallback only when the text layer is completely empty — a sparse-text
    # heuristic would mis-fire on small valid PDFs (test fixtures, short memos)
    if not text.strip():
        ocr_text = _try_vision_ocr(raw)
        if ocr_text:
            text = normalize_document_text(sanitize_input(ocr_text))
            parse_method = "pdf-ocr"
            page_count = max(page_count, 1)
        else:
            parse_method = "pdf-empty"
            warnings.append(
                {
                    "code": "scanned_no_ocr",
                    "message": "This PDF appears to be a scan — no extractable text layer.",
                }
            )

    return ParsedDocument(
        text=text,
        page_count=page_count,
        char_count=len(text),
        source_lang=detect_script(text, 4096) if text else None,
        parse_method=parse_method,
        warnings=warnings,
    )


def _parse_image(raw: bytes, mime: str) -> ParsedDocument:
    """OCR a single still image via Cloud Vision DOCUMENT_TEXT_DETECTION."""
    text = _try_vision_ocr_image(raw, mime) or ""
    text = normalize_document_text(sanitize_input(text))

    parse_method = "image-ocr" if text.strip() else "image-empty"
    warnings: list[dict] = []
    if parse_method == "image-empty":
        warnings.append(
            {
                "code": "scanned_no_ocr",
                "message": "Couldn't read text from the image. Try a clearer photo.",
            }
        )

    return ParsedDocument(
        text=text,
        page_count=1,
        char_count=len(text),
        source_lang=detect_script(text, 4096) if text else None,
        parse_method=parse_method,
        warnings=warnings,
    )


def _try_vision_ocr_image(raw: bytes, _mime: str) -> str | None:
    """Sync OCR for image bytes. Returns None on any failure (logged, not raised)."""
    try:
        from google.cloud import vision

        client = _vision_client()
        image = vision.Image(content=raw)
        resp = client.document_text_detection(image=image)
        if resp.error.message:
            logger.warning("vision image ocr error: %s", resp.error.message)
            return None
        return resp.full_text_annotation.text or None
    except Exception as exc:
        logger.warning("vision image ocr failed: %s", exc)
        return None


def _approx_page_count(text: str) -> int:
    """DOCX has no native page count; ~3000 chars per page is a reasonable rough hint."""
    return max(1, len(text) // 3000)


SUMMARIZE_SYSTEM = """You are a precise document summarizer. Output is markdown only.
Structure: ## Overview, ## Key Points (bullets), ## Action Items (numbered, if any).
Preserve named entities, dates, monetary figures verbatim. Do not invent content.
If the document is short, the Overview alone may suffice — omit empty sections."""

TRANSLATE_SYSTEM = """You are a professional translator. Translate the document below
into {target_language}. Preserve all markdown structure, headings, lists, tables,
inline emphasis, and code blocks. Keep proper nouns, brand names, file paths, and
URLs unchanged. Do not summarize, paraphrase, or omit content."""


def build_user_prompt(action: str, source_lang: str | None, parsed_text: str) -> str:
    hint = (
        f"Source language hint: {source_lang or 'unknown'}\n\n" if source_lang else ""
    )
    if action == "summarize":
        return f"{hint}Summarize the following document:\n\n---\n\n{parsed_text}"
    return f"{hint}Document to translate:\n\n---\n\n{parsed_text}"


def build_system_prompt(action: str, target_language: str) -> str:
    if action == "summarize":
        return SUMMARIZE_SYSTEM + (
            f"\n\nWrite the summary in: {target_language}." if target_language else ""
        )
    if not target_language:
        raise ValueError("target_language is required for translate")
    return TRANSLATE_SYSTEM.format(target_language=target_language)


async def stream_with_cache(
    *,
    request,
    artifact_id: UUID,
    parsed_text: str,
    source_lang: str | None,
    action: str,
    target_lang: str,
    output_cache_get,
    output_cache_set,
    history_persist,
) -> AsyncGenerator[str, None]:
    """Yields SSE-encoded JSON event lines. Caller wraps in StreamingResponse."""

    cached = await output_cache_get(artifact_id, action, target_lang)
    if cached is not None:
        yield _sse(emit_event(request, EventCode.CACHE_HIT, content=cached))
        yield _sse(emit_event(request, EventCode.DONE, tokens_used=None, model="cache"))
        return

    tokens = estimate_tokens(parsed_text)
    chunks = chunk_markdown(parsed_text, 80_000)
    yield _sse(emit_event(request, EventCode.PARSED, tokens=tokens, chunks=len(chunks)))

    from app.services.llm_service import get_llm_client

    try:
        client = get_llm_client()
    except RuntimeError as exc:
        yield _sse(
            emit_event(
                request,
                EventCode.ERROR,
                code="llm_unavailable",
                message="The assistant is unavailable right now. Try again shortly.",
            )
        )
        logger.error("llm client init failed: %s", exc)
        return

    full_prompt = build_user_prompt(action, source_lang, parsed_text)
    system = build_system_prompt(action, target_lang)

    accumulated: list[str] = []
    try:
        async for chunk in client.stream(full_prompt, system_instruction=system):
            accumulated.append(chunk)
            yield _sse(emit_event(request, EventCode.CHUNK, text=chunk))
    except Exception as exc:
        yield _sse(
            emit_event(
                request,
                EventCode.ERROR,
                code="llm_failure",
                message="Something went wrong while generating. Try again.",
            )
        )
        logger.error("llm stream failed: %s", exc)
        return

    output = "".join(accumulated)
    if not output:
        yield _sse(
            emit_event(
                request,
                EventCode.ERROR,
                code="empty_output",
                message="The assistant returned no content. Try again.",
            )
        )
        return

    await output_cache_set(
        artifact_id, action, target_lang, output, client._model_name, None
    )
    label = _history_label(action, target_lang)
    await history_persist(
        artifact_id, action, target_lang, label, output, client._model_name, None
    )

    yield _sse(
        emit_event(request, EventCode.DONE, tokens_used=None, model=client._model_name)
    )


def _sse(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


def _history_label(action: str, target_lang: str) -> str:
    verb = "Summarize" if action == "summarize" else "Translate"
    return f"{verb} → {target_lang}"
